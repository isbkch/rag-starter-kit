"""
Main document processing service that orchestrates the ingestion pipeline.
"""

import logging
import uuid
import asyncio
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime
from functools import wraps
import traceback

from app.models.documents import (
    Document,
    DocumentChunk,
    DocumentType,
    DocumentStatus,
    DocumentMetadata,
)
from app.services.ingestion.text_extractor import TextExtractor
from app.services.ingestion.chunk_processor import ChunkProcessor, ChunkingConfig
from app.services.search.embedding_service import get_embedding_service
from app.services.vectordb.factory import VectorDBFactory
from app.core.config import settings
from app.core.tracing import trace_document_processing, get_metrics_collector

logger = logging.getLogger(__name__)


def retry_on_failure(max_retries: int = 3, delay: float = 1.0, backoff: float = 2.0):
    """Decorator to retry failed operations with exponential backoff."""

    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            last_exception = None

            for attempt in range(max_retries + 1):
                try:
                    if asyncio.iscoroutinefunction(func):
                        return await func(*args, **kwargs)
                    else:
                        return func(*args, **kwargs)
                except Exception as e:
                    last_exception = e

                    if attempt == max_retries:
                        logger.error(f"Final attempt failed for {func.__name__}: {e}")
                        break

                    wait_time = delay * (backoff**attempt)
                    logger.warning(
                        f"Attempt {attempt + 1} failed for {func.__name__}, retrying in {wait_time:.1f}s: {e}"
                    )
                    await asyncio.sleep(wait_time)

            raise last_exception

        return wrapper

    return decorator


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""

    def __init__(
        self, message: str, error_code: str = None, details: Dict[str, Any] = None
    ):
        super().__init__(message)
        self.error_code = error_code or "PROCESSING_ERROR"
        self.details = details or {}
        self.timestamp = datetime.utcnow()


class ValidationError(DocumentProcessingError):
    """Validation error."""

    def __init__(self, message: str, details: Dict[str, Any] = None):
        super().__init__(message, "VALIDATION_ERROR", details)


class ExtractionError(DocumentProcessingError):
    """Text extraction error."""

    def __init__(self, message: str, details: Dict[str, Any] = None):
        super().__init__(message, "EXTRACTION_ERROR", details)


class ChunkingError(DocumentProcessingError):
    """Chunking error."""

    def __init__(self, message: str, details: Dict[str, Any] = None):
        super().__init__(message, "CHUNKING_ERROR", details)


class EmbeddingError(DocumentProcessingError):
    """Embedding generation error."""

    def __init__(self, message: str, details: Dict[str, Any] = None):
        super().__init__(message, "EMBEDDING_ERROR", details)


class VectorStorageError(DocumentProcessingError):
    """Vector storage error."""

    def __init__(self, message: str, details: Dict[str, Any] = None):
        super().__init__(message, "VECTOR_STORAGE_ERROR", details)


class DocumentProcessor:
    """Main document processing service."""

    def __init__(self):
        self.text_extractor = TextExtractor()
        self.chunk_processor = ChunkProcessor()

    @trace_document_processing(operation="process", filename="")
    async def process_document(
        self,
        file_path: str,
        original_filename: str,
        document_type: Optional[DocumentType] = None,
        chunking_strategy: str = "recursive",
        chunking_config: Optional[ChunkingConfig] = None,
        extract_metadata: bool = True,
        generate_embeddings: bool = True,
        store_in_vector_db: bool = True,
    ) -> Document:
        """
        Process a document through the complete ingestion pipeline.

        Args:
            file_path: Path to the document file
            original_filename: Original filename
            document_type: Document type (auto-detected if None)
            chunking_strategy: Strategy for text chunking
            chunking_config: Configuration for chunking
            extract_metadata: Whether to extract metadata
            generate_embeddings: Whether to generate embeddings for chunks
            store_in_vector_db: Whether to store chunks in vector database

        Returns:
            Processed Document object
        """
        import time

        start_time = time.time()
        metrics = get_metrics_collector()

        document_id = str(uuid.uuid4())
        document = None

        try:
            # Auto-detect document type if not provided
            if document_type is None:
                document_type = self.text_extractor.detect_document_type(file_path)

            # Create initial document object
            document = Document(
                id=document_id,
                filename=Path(file_path).name,
                original_filename=original_filename,
                file_path=file_path,
                document_type=document_type,
                status=DocumentStatus.PROCESSING,
                metadata=DocumentMetadata(),
                chunks=[],
                total_chunks=0,
                processed_chunks=0,
            )

            logger.info(
                f"Starting processing for document {document_id}: {original_filename}"
            )

            # Step 1: Extract text
            text_content = await self._extract_text(document)

            # Step 2: Extract metadata if requested
            if extract_metadata:
                document.metadata = await self._extract_metadata(document)

            # Step 3: Process chunks
            chunks = await self._process_chunks(
                document_id=document_id,
                text=text_content,
                chunking_strategy=chunking_strategy,
                chunking_config=chunking_config,
            )

            # Step 4: Generate embeddings if requested
            if generate_embeddings:
                chunks = await self._generate_embeddings(chunks)

            # Step 5: Store in vector database if requested
            if store_in_vector_db and generate_embeddings:
                await self._store_in_vector_db(chunks)

            # Step 6: Update document with results
            document.chunks = chunks
            document.total_chunks = len(chunks)
            document.processed_chunks = len(chunks)
            document.status = DocumentStatus.COMPLETED
            document.updated_at = datetime.utcnow()

            logger.info(
                f"Successfully processed document {document_id} with {len(chunks)} chunks"
            )

            # Record successful processing metrics
            duration = time.time() - start_time
            metrics.record_document_processing("process", duration, success=True)

            return document

        except DocumentProcessingError as e:
            logger.error(f"Document processing error for {document_id}: {e}")

            # Update document with error status if we have a document object
            if document:
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
                document.error_code = e.error_code
                document.error_details = e.details
                document.updated_at = datetime.utcnow()

            raise
        except Exception as e:
            logger.error(f"Unexpected error processing document {document_id}: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")

            # Update document with error status if we have a document object
            if document:
                document.status = DocumentStatus.FAILED
                document.error_message = f"Unexpected error: {str(e)}"
                document.error_code = "UNEXPECTED_ERROR"
                document.error_details = {"traceback": traceback.format_exc()}
                document.updated_at = datetime.utcnow()

            # Record failed processing metrics
            duration = time.time() - start_time
            metrics.record_document_processing("process", duration, success=False)
            metrics.record_error(type(e).__name__, "document.process")

            raise DocumentProcessingError(
                f"Unexpected error processing document: {e}",
                "UNEXPECTED_ERROR",
                {"original_error": str(e), "traceback": traceback.format_exc()},
            )

    @retry_on_failure(max_retries=2, delay=0.5)
    async def _extract_text(self, document: Document) -> str:
        """Extract text from document with retry logic."""
        try:
            # Validate file before extraction
            if not Path(document.file_path).exists():
                raise ExtractionError(
                    f"File not found: {document.file_path}",
                    {"file_path": document.file_path, "document_id": document.id},
                )

            text = self.text_extractor.extract_text(
                document.file_path, document.document_type
            )

            if not text or not text.strip():
                raise ExtractionError(
                    "No text content extracted from document",
                    {
                        "file_path": document.file_path,
                        "document_type": document.document_type.value,
                        "file_size": Path(document.file_path).stat().st_size,
                    },
                )

            # Validate extracted text
            text_length = len(text)
            if text_length < 10:
                logger.warning(
                    f"Very short text extracted ({text_length} chars) from {document.filename}"
                )

            logger.debug(f"Extracted {text_length} characters from {document.filename}")
            return text

        except ExtractionError:
            raise
        except Exception as e:
            logger.error(f"Text extraction failed for {document.filename}: {e}")
            raise ExtractionError(
                f"Failed to extract text: {e}",
                {
                    "file_path": document.file_path,
                    "document_type": document.document_type.value,
                    "original_error": str(e),
                },
            )

    async def _extract_metadata(self, document: Document) -> DocumentMetadata:
        """Extract metadata from document."""
        try:
            metadata = self.text_extractor.extract_metadata(
                document.file_path, document.document_type
            )

            # Add processing metadata
            metadata.custom_fields.update(
                {
                    "processing_timestamp": datetime.utcnow().isoformat(),
                    "processor_version": "1.0.0",
                    "document_id": document.id,
                }
            )

            logger.debug(f"Extracted metadata for {document.filename}")
            return metadata

        except Exception as e:
            logger.warning(f"Metadata extraction failed for {document.filename}: {e}")
            # Return basic metadata on failure
            return DocumentMetadata(
                file_size=Path(document.file_path).stat().st_size,
                custom_fields={
                    "metadata_extraction_error": str(e),
                    "processing_timestamp": datetime.utcnow().isoformat(),
                },
            )

    async def _process_chunks(
        self,
        document_id: str,
        text: str,
        chunking_strategy: str,
        chunking_config: Optional[ChunkingConfig],
    ) -> List[DocumentChunk]:
        """Process text into chunks with validation."""
        try:
            # Validate input text
            if not text or len(text.strip()) < 10:
                raise ChunkingError(
                    "Text too short for chunking",
                    {"text_length": len(text), "document_id": document_id},
                )

            # Optimize chunking config if not provided
            if chunking_config is None:
                chunking_config = self.chunk_processor.optimize_chunking_config(text)

            chunks = self.chunk_processor.process_document(
                document_id=document_id,
                text=text,
                chunking_strategy=chunking_strategy,
                config=chunking_config,
            )

            # Validate chunks
            if not chunks:
                raise ChunkingError(
                    "No chunks generated from text",
                    {
                        "text_length": len(text),
                        "chunking_strategy": chunking_strategy,
                        "document_id": document_id,
                    },
                )

            # Validate chunk quality
            valid_chunks = []
            for chunk in chunks:
                if chunk.content and len(chunk.content.strip()) >= 5:
                    valid_chunks.append(chunk)
                else:
                    logger.warning(
                        f"Skipping invalid chunk {chunk.id} (too short or empty)"
                    )

            if not valid_chunks:
                raise ChunkingError(
                    "No valid chunks generated",
                    {"total_chunks": len(chunks), "document_id": document_id},
                )

            # Log chunk statistics
            stats = self.chunk_processor.get_chunk_statistics(valid_chunks)
            logger.info(f"Chunk statistics for {document_id}: {stats}")

            return valid_chunks

        except ChunkingError:
            raise
        except Exception as e:
            logger.error(f"Chunk processing failed for {document_id}: {e}")
            raise ChunkingError(
                f"Failed to process chunks: {e}",
                {"document_id": document_id, "original_error": str(e)},
            )

    @retry_on_failure(max_retries=3, delay=1.0)
    async def _generate_embeddings(
        self, chunks: List[DocumentChunk]
    ) -> List[DocumentChunk]:
        """Generate embeddings for document chunks with retry logic."""
        try:
            if not chunks:
                raise EmbeddingError("No chunks provided for embedding generation")

            # Get embedding service
            embedding_service = await get_embedding_service()

            # Validate chunks before processing
            valid_chunks = []
            chunk_texts = []

            for chunk in chunks:
                if chunk.content and len(chunk.content.strip()) > 0:
                    valid_chunks.append(chunk)
                    chunk_texts.append(chunk.content)
                else:
                    logger.warning(f"Skipping chunk {chunk.id} with empty content")

            if not chunk_texts:
                raise EmbeddingError("No valid chunks with content found")

            logger.info(f"Generating embeddings for {len(chunk_texts)} chunks")

            # Generate embeddings in smaller batches for better error handling
            batch_size = min(50, len(chunk_texts))
            all_embeddings = []

            for i in range(0, len(chunk_texts), batch_size):
                batch_texts = chunk_texts[i : i + batch_size]
                logger.debug(
                    f"Processing embedding batch {i // batch_size + 1}/{(len(chunk_texts) + batch_size - 1) // batch_size}"
                )

                try:
                    batch_embeddings = await embedding_service.get_embeddings(
                        batch_texts
                    )
                    if not batch_embeddings or len(batch_embeddings) != len(
                        batch_texts
                    ):
                        raise EmbeddingError(
                            f"Embedding batch returned incorrect number of embeddings",
                            {
                                "expected": len(batch_texts),
                                "received": len(batch_embeddings)
                                if batch_embeddings
                                else 0,
                                "batch_index": i // batch_size,
                            },
                        )
                    all_embeddings.extend(batch_embeddings)
                except Exception as e:
                    raise EmbeddingError(
                        f"Failed to generate embeddings for batch {i // batch_size + 1}: {e}",
                        {
                            "batch_index": i // batch_size,
                            "batch_size": len(batch_texts),
                        },
                    )

            # Assign embeddings to chunks
            successful_chunks = []
            for chunk, embedding in zip(valid_chunks, all_embeddings):
                if embedding and len(embedding) > 0:
                    chunk.embedding = embedding
                    # Add embedding metadata
                    chunk.metadata.update(
                        {
                            "embedding_model": embedding_service.model_name,
                            "embedding_dimension": len(embedding),
                            "embedding_generated_at": datetime.utcnow().isoformat(),
                        }
                    )
                    successful_chunks.append(chunk)
                else:
                    logger.warning(f"Chunk {chunk.id} received empty embedding")

            if not successful_chunks:
                raise EmbeddingError("No chunks received valid embeddings")

            logger.info(
                f"Successfully generated embeddings for {len(successful_chunks)} chunks"
            )
            return successful_chunks

        except EmbeddingError:
            raise
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            raise EmbeddingError(
                f"Unexpected error generating embeddings: {e}",
                {"original_error": str(e)},
            )

    @retry_on_failure(max_retries=3, delay=2.0)
    async def _store_in_vector_db(self, chunks: List[DocumentChunk]) -> None:
        """Store chunks in vector database with retry logic."""
        vector_db = None
        try:
            if not chunks:
                raise VectorStorageError("No chunks provided for storage")

            # Validate chunks have embeddings
            chunks_with_embeddings = []
            for chunk in chunks:
                if chunk.embedding and len(chunk.embedding) > 0:
                    chunks_with_embeddings.append(chunk)
                else:
                    logger.warning(
                        f"Chunk {chunk.id} missing embedding, skipping storage"
                    )

            if not chunks_with_embeddings:
                raise VectorStorageError("No chunks with valid embeddings found")

            # Get vector database instance
            vector_db = VectorDBFactory.create_vector_db(
                provider=settings.VECTOR_DB_PROVIDER
            )

            # Test connection first
            await vector_db.connect()
            health = await vector_db.health_check()
            if health["status"] != "healthy":
                raise VectorStorageError(
                    f"Vector database is not healthy: {health.get('error', 'Unknown error')}",
                    {"health_status": health},
                )

            logger.info(
                f"Storing {len(chunks_with_embeddings)} chunks in vector database"
            )

            # Store chunks using the base class helper method
            chunk_ids = await vector_db.insert_document_chunks(chunks_with_embeddings)

            if not chunk_ids or len(chunk_ids) != len(chunks_with_embeddings):
                raise VectorStorageError(
                    "Chunk storage returned unexpected results",
                    {
                        "expected_chunks": len(chunks_with_embeddings),
                        "stored_chunks": len(chunk_ids) if chunk_ids else 0,
                    },
                )

            logger.info(
                f"Successfully stored {len(chunk_ids)} chunks in vector database"
            )

        except VectorStorageError:
            raise
        except Exception as e:
            logger.error(f"Failed to store chunks in vector database: {e}")
            raise VectorStorageError(
                f"Unexpected error storing chunks: {e}", {"original_error": str(e)}
            )
        finally:
            # Always try to disconnect
            if vector_db:
                try:
                    await vector_db.disconnect()
                except Exception as e:
                    logger.warning(f"Error disconnecting from vector database: {e}")

    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """
        Validate document before processing with comprehensive checks.

        Returns:
            Dictionary with validation results
        """
        validation_result = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "file_info": {},
        }

        try:
            file_path_obj = Path(file_path)

            # Check if file exists
            if not file_path_obj.exists():
                validation_result["is_valid"] = False
                validation_result["errors"].append("File does not exist")
                return validation_result

            # Check if it's actually a file
            if not file_path_obj.is_file():
                validation_result["is_valid"] = False
                validation_result["errors"].append("Path is not a file")
                return validation_result

            # Check file size
            file_size = file_path_obj.stat().st_size
            validation_result["file_info"]["size"] = file_size
            validation_result["file_info"]["size_mb"] = round(
                file_size / (1024 * 1024), 2
            )

            if file_size > settings.MAX_FILE_SIZE:
                validation_result["is_valid"] = False
                validation_result["errors"].append(
                    f"File size ({file_size} bytes) exceeds maximum allowed size ({settings.MAX_FILE_SIZE} bytes)"
                )

            if file_size == 0:
                validation_result["is_valid"] = False
                validation_result["errors"].append("File is empty")

            # Warn about very large files
            if file_size > settings.MAX_FILE_SIZE * 0.8:
                validation_result["warnings"].append(
                    f'File is quite large ({validation_result["file_info"]["size_mb"]} MB), processing may take time'
                )

            # Check file extension and type
            try:
                document_type = self.text_extractor.detect_document_type(file_path)
                validation_result["file_info"]["document_type"] = document_type.value
                validation_result["file_info"][
                    "file_extension"
                ] = file_path_obj.suffix.lower()
            except Exception as e:
                validation_result["is_valid"] = False
                validation_result["errors"].append(f"Unsupported file type: {e}")
                return validation_result

            # Check if file is readable
            try:
                with open(file_path, "rb") as f:
                    header = f.read(1024)  # Try to read first 1KB
                    if not header:
                        validation_result["errors"].append(
                            "File appears to be empty or unreadable"
                        )
                        validation_result["is_valid"] = False
            except PermissionError:
                validation_result["is_valid"] = False
                validation_result["errors"].append("Permission denied reading file")
            except Exception as e:
                validation_result["is_valid"] = False
                validation_result["errors"].append(f"File is not readable: {e}")

            # Document type specific validation
            if document_type == DocumentType.PDF:
                pdf_validation = self._validate_pdf(file_path)
                validation_result["errors"].extend(pdf_validation.get("errors", []))
                validation_result["warnings"].extend(pdf_validation.get("warnings", []))
                if pdf_validation.get("errors"):
                    validation_result["is_valid"] = False
            elif document_type == DocumentType.DOCX:
                docx_validation = self._validate_docx(file_path)
                validation_result["errors"].extend(docx_validation.get("errors", []))
                validation_result["warnings"].extend(
                    docx_validation.get("warnings", [])
                )
                if docx_validation.get("errors"):
                    validation_result["is_valid"] = False

        except Exception as e:
            validation_result["is_valid"] = False
            validation_result["errors"].append(f"Validation error: {e}")
            logger.error(f"Document validation error for {file_path}: {e}")

        return validation_result

    def _validate_pdf(self, file_path: str) -> Dict[str, Any]:
        """Validate PDF file."""
        result = {"errors": [], "warnings": []}

        try:
            import PyPDF2

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    result["errors"].append("PDF is encrypted and cannot be processed")

                # Check number of pages
                page_count = len(pdf_reader.pages)
                if page_count == 0:
                    result["errors"].append("PDF has no pages")
                elif page_count > 1000:
                    result["warnings"].append(
                        f"PDF has many pages ({page_count}), processing may take time"
                    )

        except Exception as e:
            result["errors"].append(f"PDF validation error: {e}")

        return result

    def _validate_docx(self, file_path: str) -> Dict[str, Any]:
        """Validate DOCX file."""
        result = {"errors": [], "warnings": []}

        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)

            # Check if document has content
            if not doc.paragraphs:
                result["warnings"].append("DOCX appears to have no paragraphs")

        except Exception as e:
            result["errors"].append(f"DOCX validation error: {e}")

        return result

    def get_supported_formats(self) -> List[Dict[str, Any]]:
        """Get list of supported document formats."""
        return [
            {
                "type": DocumentType.PDF.value,
                "extensions": [".pdf"],
                "description": "Adobe PDF documents",
                "max_size_mb": settings.MAX_FILE_SIZE // (1024 * 1024),
            },
            {
                "type": DocumentType.DOCX.value,
                "extensions": [".docx"],
                "description": "Microsoft Word documents",
                "max_size_mb": settings.MAX_FILE_SIZE // (1024 * 1024),
            },
            {
                "type": DocumentType.MARKDOWN.value,
                "extensions": [".md", ".markdown"],
                "description": "Markdown documents",
                "max_size_mb": settings.MAX_FILE_SIZE // (1024 * 1024),
            },
            {
                "type": DocumentType.TXT.value,
                "extensions": [".txt"],
                "description": "Plain text documents",
                "max_size_mb": settings.MAX_FILE_SIZE // (1024 * 1024),
            },
        ]

    def estimate_processing_time(self, file_path: str) -> Dict[str, Any]:
        """Estimate processing time for a document."""
        try:
            file_size = Path(file_path).stat().st_size
            document_type = self.text_extractor.detect_document_type(file_path)

            # Rough estimates based on file size and type
            base_time = 0
            if document_type == DocumentType.PDF:
                base_time = file_size / (1024 * 1024) * 5  # 5 seconds per MB for PDF
            elif document_type == DocumentType.DOCX:
                base_time = file_size / (1024 * 1024) * 2  # 2 seconds per MB for DOCX
            else:
                base_time = (
                    file_size / (1024 * 1024) * 1
                )  # 1 second per MB for text files

            return {
                "estimated_seconds": max(1, int(base_time)),
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "document_type": document_type.value,
            }

        except Exception as e:
            return {
                "estimated_seconds": 30,  # Default estimate
                "error": str(e),
            }
