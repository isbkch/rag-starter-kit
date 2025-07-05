"""
Text extraction service for different document types.
"""

import io
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from abc import ABC, abstractmethod

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

from app.models.documents import DocumentType, DocumentMetadata

logger = logging.getLogger(__name__)


class BaseTextExtractor(ABC):
    """Base class for text extractors."""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from file."""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from file."""
        pass


class PDFTextExtractor(BaseTextExtractor):
    """PDF text extractor using multiple libraries for robustness."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            # Try pdfplumber first (better for complex layouts)
            return self._extract_with_pdfplumber(file_path)
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            try:
                # Fallback to PyPDF2
                return self._extract_with_pypdf2(file_path)
            except Exception as e2:
                logger.error(f"Both PDF extractors failed for {file_path}: {e2}")
                raise ValueError(f"Failed to extract text from PDF: {e2}")
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber."""
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n\n".join(text_content)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2."""
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n\n".join(text_content)
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from PDF file."""
        metadata = DocumentMetadata()
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Basic metadata
                metadata.page_count = len(pdf_reader.pages)
                metadata.file_size = Path(file_path).stat().st_size
                
                # PDF metadata
                if pdf_reader.metadata:
                    metadata.title = pdf_reader.metadata.get('/Title')
                    metadata.author = pdf_reader.metadata.get('/Author')
                    
                    # Convert PDF dates to datetime
                    created_date = pdf_reader.metadata.get('/CreationDate')
                    if created_date:
                        metadata.created_date = self._parse_pdf_date(created_date)
                    
                    modified_date = pdf_reader.metadata.get('/ModDate')
                    if modified_date:
                        metadata.modified_date = self._parse_pdf_date(modified_date)
                        
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata from {file_path}: {e}")
        
        return metadata
    
    def _parse_pdf_date(self, pdf_date: str) -> Optional[str]:
        """Parse PDF date format to ISO format."""
        try:
            # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
            if pdf_date.startswith('D:'):
                date_str = pdf_date[2:16]  # Extract YYYYMMDDHHMMSS
                return f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}T{date_str[8:10]}:{date_str[10:12]}:{date_str[12:14]}"
        except Exception:
            pass
        return None


class DOCXTextExtractor(BaseTextExtractor):
    """DOCX text extractor."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from DOCX file."""
        metadata = DocumentMetadata()
        
        try:
            doc = DocxDocument(file_path)
            metadata.file_size = Path(file_path).stat().st_size
            
            # Core properties
            if doc.core_properties:
                metadata.title = doc.core_properties.title
                metadata.author = doc.core_properties.author
                metadata.created_date = doc.core_properties.created
                metadata.modified_date = doc.core_properties.modified
                
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata from {file_path}: {e}")
        
        return metadata


class MarkdownTextExtractor(BaseTextExtractor):
    """Markdown text extractor."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                markdown_content = file.read()
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(markdown_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            return soup.get_text()
        except Exception as e:
            logger.error(f"Failed to extract text from Markdown {file_path}: {e}")
            raise ValueError(f"Failed to extract text from Markdown: {e}")
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from Markdown file."""
        metadata = DocumentMetadata()
        
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            metadata.file_size = stat.st_size
            metadata.created_date = stat.st_ctime
            metadata.modified_date = stat.st_mtime
            
            # Extract title from first heading
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                lines = content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        metadata.title = line[2:].strip()
                        break
                        
        except Exception as e:
            logger.warning(f"Failed to extract Markdown metadata from {file_path}: {e}")
        
        return metadata


class TextFileExtractor(BaseTextExtractor):
    """Plain text file extractor."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Failed to extract text from {file_path}: {e}")
                raise ValueError(f"Failed to extract text from file: {e}")
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from text file."""
        metadata = DocumentMetadata()
        
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            metadata.file_size = stat.st_size
            metadata.created_date = stat.st_ctime
            metadata.modified_date = stat.st_mtime
            
        except Exception as e:
            logger.warning(f"Failed to extract text metadata from {file_path}: {e}")
        
        return metadata


class TextExtractor:
    """Main text extraction service."""
    
    def __init__(self):
        self.extractors = {
            DocumentType.PDF: PDFTextExtractor(),
            DocumentType.DOCX: DOCXTextExtractor(),
            DocumentType.MARKDOWN: MarkdownTextExtractor(),
            DocumentType.TXT: TextFileExtractor(),
        }
    
    def extract_text(self, file_path: str, document_type: DocumentType) -> str:
        """Extract text from file based on document type."""
        extractor = self.extractors.get(document_type)
        if not extractor:
            raise ValueError(f"Unsupported document type: {document_type}")
        
        return extractor.extract_text(file_path)
    
    def extract_metadata(self, file_path: str, document_type: DocumentType) -> DocumentMetadata:
        """Extract metadata from file based on document type."""
        extractor = self.extractors.get(document_type)
        if not extractor:
            raise ValueError(f"Unsupported document type: {document_type}")
        
        return extractor.extract_metadata(file_path)
    
    def detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension."""
        file_path_obj = Path(file_path)
        extension = file_path_obj.suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.txt': DocumentType.TXT,
        }
        
        return type_mapping.get(extension, DocumentType.TXT) 